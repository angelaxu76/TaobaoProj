import csv
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# finance_config.py 与本文件在同一目录
from finance_config import FINANCE_EES


# ✅ 不向香港公司收费的类别（英国自己承担的 or 不相关的）
EXCLUDE_CATEGORIES = {
    "Sales",                 # 香港打款，不是成本/收入
    "Accountant",            # 会计费用：英国公司承担
    "Business account fees", # ANNA 账户费：英国公司承担
}



def load_and_prepare(csv_path: str) -> pd.DataFrame:
    """
    读取 ANNA 导出的 CSV，做基础清洗，并返回 DataFrame.

    关键点：
    - 容忍坏行（列数不对），自动截断/补齐
    - 转换 Amount 为 float
    - 解析 Created 为 datetime
    - 添加 Item_Type 分类列
    - 过滤掉 EXCLUDE_CATEGORIES（不向香港收费）
    - 添加 Net_Ex_VAT / VAT_Amount（按 20% 反推）
    """
    csv_path = Path(csv_path)
    print(f"[INFO] Loading CSV: {csv_path}")

    rows = []
    with csv_path.open("r", encoding="utf-8", newline="") as f:
        reader = csv.reader(f)
        try:
            header = next(reader)
        except StopIteration:
            raise ValueError("Empty CSV")

        expected_cols = len(header)

        for line_no, row in enumerate(reader, start=2):
            if not any(row):
                # 空行跳过
                continue

            if len(row) < expected_cols:
                # 列数不够 -> 用空字符串补齐
                row = row + [""] * (expected_cols - len(row))
            elif len(row) > expected_cols:
                # 列数太多 -> 截断
                row = row[:expected_cols]

            rows.append(row)

    df = pd.DataFrame(rows, columns=header)
    # 去掉全空列
    df = df.dropna(axis=1, how="all")

    # ---- 解析日期 ----
    if "Created" in df.columns:
        # ANNA 一般是 "2025-11-02 10:23:45"
        df["Created_dt"] = pd.to_datetime(
            df["Created"].str.replace(",", ""),
            format="%Y-%m-%d %H:%M:%S",
            errors="coerce",
        )
    else:
        df["Created_dt"] = pd.NaT

    # ---- 金额转 float ----
    if "Amount" not in df.columns:
        raise ValueError("CSV 中缺少 Amount 列")

    df["Amount"] = pd.to_numeric(df["Amount"], errors="coerce").fillna(0.0)

    # ---- 分类（仅用于展示）----
    def classify_item_type(category: str, description: str) -> str:
        category = (category or "").lower()
        description = (description or "").lower()

        if "stock" in category or "inventory" in category:
            return "Stock / Goods"
        if "postage" in category or "shipping" in category or "delivery" in category:
            return "Shipping / Postage"
        if "materials" in category or "packag" in category:
            return "Packaging / Materials"
        if "parcel" in category or "courier" in description:
            return "Courier / Parcel"
        if "refund" in category or "return" in description:
            return "Refund / Return"
        if "fee" in category:
            return "Fees / Charges"
        return "Other cost"

    df["Item_Type"] = df.apply(
        lambda r: classify_item_type(
            r.get("Category", "") if "Category" in df.columns else "",
            r.get("Description", "") if "Description" in df.columns else "",
        ),
        axis=1,
    )

    # ---- 过滤掉不向香港收费的类别 ----
    if "Category" in df.columns:
        before = len(df)
        df = df[~df["Category"].isin(EXCLUDE_CATEGORIES)].copy()
        after = len(df)
        print(f"[INFO] Filter EXCLUDE_CATEGORIES: {before} -> {after} rows")

    # ---- 计算 Net_Ex_VAT / VAT_Amount（按 20% 反推） ----
    # 约定：Amount 为负数代表支出，为正数代表退款
    # 所以直接 Amount / 1.2 即为“含符号的净额”
    df["Net_Ex_VAT"] = (df["Amount"] / 1.2).round(2)
    df["VAT_Amount"] = (df["Amount"] - df["Net_Ex_VAT"]).round(2)

    return df


def create_accounting_report(df: pd.DataFrame, output_path: Path) -> Path:
    """
    创建一个 Excel 记账报表：
    - Sheet1: 明细（含 Amount, Net_Ex_VAT, VAT_Amount, Item_Type）
    - Sheet2: 汇总（按 Item_Type 分组 + 总计）
    - 再附上 Net total / VAT total / Gross total / UK Profit / HK Payable
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    print(f"[INFO] Writing accounting report to: {output_path}")

    # 只保留一些对你记账有意义的字段
    detail_cols = []
    for col in [
        "Created",
        "Created_dt",
        "Description",
        "Category",
        "Amount",
        "Net_Ex_VAT",
        "VAT_Amount",
        "Item_Type",
        "Reference",
        "Counterparty name",
    ]:
        if col in df.columns:
            detail_cols.append(col)

    df_detail = df[detail_cols].copy()

    # 计算总额
    gross_total = df["Amount"].sum().round(2)          # 含 VAT（带符号）
    net_total = df["Net_Ex_VAT"].sum().round(2)        # 去 VAT（带符号）
    vat_total = df["VAT_Amount"].sum().round(2)

    # 取“正向成本数值”
    gross_cost = round(-gross_total, 2)  # 负数 → 成本
    net_cost = round(-net_total, 2)

    uk_profit = round(net_cost * 0.15, 2)
    hk_payable = round(net_cost * 1.15, 2)

    # 分组汇总
    grp = df.groupby("Item_Type")[["Amount", "Net_Ex_VAT", "VAT_Amount"]].sum().round(2)
    grp = grp.reset_index()

    # 写 Excel
    with pd.ExcelWriter(output_path, engine="xlsxwriter") as writer:
        df_detail.to_excel(writer, sheet_name="Details", index=False)

        summary_rows = []

        for _, row in grp.iterrows():
            summary_rows.append(
                {
                    "Item_Type": row["Item_Type"],
                    "Amount (Gross)": row["Amount"],
                    "Net_Ex_VAT": row["Net_Ex_VAT"],
                    "VAT_Amount": row["VAT_Amount"],
                }
            )

        summary_rows.append({})
        summary_rows.append({"Item_Type": "TOTAL (signed)", "Amount (Gross)": gross_total,
                             "Net_Ex_VAT": net_total, "VAT_Amount": vat_total})
        summary_rows.append({"Item_Type": "COST (positive)", "Amount (Gross)": gross_cost,
                             "Net_Ex_VAT": net_cost, "VAT_Amount": vat_total})
        summary_rows.append({"Item_Type": "UK Profit (15%)", "Amount (Gross)": uk_profit})
        summary_rows.append({"Item_Type": "HK Payable (Net*1.15)", "Amount (Gross)": hk_payable})

        df_summary = pd.DataFrame(summary_rows)
        df_summary.to_excel(writer, sheet_name="Summary", index=False)

    print(f"[OK] Accounting report: {output_path}")
    return output_path


def _set_paragraph_style(p, bold=False, size=11, align=None):
    """小工具：设置段落基本样式。"""
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
    if align is not None:
        p.alignment = align


def create_invoice_docx(
    df: pd.DataFrame,
    csv_path: Path,
    output_path: Path,
) -> Path:
    """
    根据 df 生成一个正式的 Intercompany Invoice（DOCX，英文商务版）。

    金额逻辑：
    - gross_total = df["Amount"].sum()        （带符号，负数为成本）
    - net_total   = df["Net_Ex_VAT"].sum()
    - gross_cost  = -gross_total             （正数：含 VAT 成本）
    - net_cost    = -net_total               （正数：去 VAT 成本 = 香港成本基数）
    - uk_profit   = net_cost * 0.15
    - hk_payable  = net_cost * 1.15          （发票金额）
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    print(f"[INFO] Writing invoice DOCX to: {output_path}")

    gross_total = df["Amount"].sum().round(2)
    net_total = df["Net_Ex_VAT"].sum().round(2)
    vat_total = df["VAT_Amount"].sum().round(2)

    gross_cost = round(-gross_total, 2)
    net_cost = round(-net_total, 2)
    uk_profit = round(net_cost * 0.15, 2)
    hk_payable = round(net_cost * 1.15, 2)

    # 取账期：按 Created_dt 的最小/最大日期
    if df["Created_dt"].notna().any():
        start_date = df["Created_dt"].min().date()
        end_date = df["Created_dt"].max().date()
    else:
        # fallback：从文件名推一个月份
        try:
            # e.g. 2025-11.csv
            stem = Path(csv_path).stem
            dt = datetime.strptime(stem, "%Y-%m")
            start_date = dt.date().replace(day=1)
            end_date = dt.date()
        except Exception:
            start_date = None
            end_date = None

    exporter = FINANCE_EES["exporter"]
    consignee = FINANCE_EES["consignee"]

    doc = Document()

    # 标题
    title = doc.add_heading("INVOICE", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息表头
    p = doc.add_paragraph()
    run = p.add_run(f"Invoice Date: {datetime.today().date().isoformat()}")
    _set_paragraph_style(p, size=11)

    period_text = ""
    if start_date and end_date:
        period_text = f"{start_date.isoformat()} to {end_date.isoformat()}"
    else:
        period_text = "See attached ANNA statement"

    p = doc.add_paragraph()
    run = p.add_run(f"Period Covered: {period_text}")
    _set_paragraph_style(p, size=11)

    doc.add_paragraph()  # 空行

    # 出票方（英国公司）
    p = doc.add_paragraph()
    run = p.add_run("From (Supplier):\n")
    _set_paragraph_style(p, bold=True, size=11)
    p.add_run(f"{exporter['name']}\n")
    p.add_run(f"Company No: {exporter.get('company_no', '')}\n")
    p.add_run(f"VAT No: {exporter.get('vat_no', '')}\n")
    p.add_run(f"EORI: {exporter.get('eori_no', '')}\n")
    p.add_run(f"Address: {exporter['address']}\n")
    p.add_run(f"Email: {exporter.get('email', '')}\n")
    p.add_run(f"Phone: {exporter.get('phone', '')}")

    doc.add_paragraph()  # 空行

    # 收票方（香港公司）
    p = doc.add_paragraph()
    run = p.add_run("Bill To (Customer):\n")
    _set_paragraph_style(p, bold=True, size=11)
    p.add_run(f"{consignee['name']}\n")
    p.add_run(f"Address: {consignee['address']}\n")
    p.add_run(f"Email: {consignee.get('email', '')}\n")
    p.add_run(f"Phone: {consignee.get('phone', '')}")

    doc.add_paragraph()  # 空行

    # 定价说明
    p = doc.add_paragraph()
    run = p.add_run("Pricing Model:\n")
    _set_paragraph_style(p, bold=True, size=11)
    p.add_run(
        "The Hong Kong company reimburses the net purchase cost of UK sourcing "
        "(excluding UK VAT) plus a 15% resale margin. All VAT input tax is "
        "claimed and retained by the UK company and is not charged to the Hong Kong company."
    )

    doc.add_paragraph()  # 空行

    # 金额表格
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Description"
    hdr_cells[1].text = "Amount (GBP)"

    def add_row(label: str, value: float):
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = f"{value:,.2f}"

    add_row("Total ANNA cost (gross, incl. VAT, signed)", gross_total)
    add_row("Total ANNA cost (net of VAT, signed)", net_total)
    add_row("Total VAT amount (signed)", vat_total)
    add_row("Net purchase cost (positive)", net_cost)
    add_row("UK resale margin (15% of net cost)", uk_profit)
    add_row("Total payable by HK company (net × 1.15)", hk_payable)

    doc.add_paragraph()  # 空行

    # 备注
    p = doc.add_paragraph()
    p.add_run(
        "Notes:\n"
        "- The above invoice covers all ANNA bank transactions related to stock purchase, "
        "shipping, packaging and other direct sourcing costs for the specified period, "
        "excluding UK-only overheads such as accountant fees and ANNA business account fees.\n"
        "- The Hong Kong company bears the net purchase cost and logistics related costs "
        "only. UK VAT input tax is fully retained by the UK company.\n"
        "- This invoice is issued under the intercompany trade agreement and pricing "
        "policy between EMINZORA TRADE LTD and HONG KONG ANGEL XUAN TRADING CO., LIMITED."
    )
    _set_paragraph_style(p, size=10)

    doc.add_paragraph()  # 空行

    # 签名区
    p = doc.add_paragraph("\n\nAuthorised Signature (UK): ____________________________")
    _set_paragraph_style(p, size=11)

    doc.save(output_path)
    print(f"[OK] Invoice DOCX: {output_path}")
    return output_path


def generate_anna_monthly_reports(csv_path: str, output_dir: str):
    """
    主入口函数（供 pipeline 调用）：

    输入：
        csv_path: ANNA 月度交易 CSV 路径
        output_dir: 输出目录

    输出：
        (accounting_path, invoice_path)
    """
    csv_path = Path(csv_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    df = load_and_prepare(str(csv_path))

    # 生成记账报表
    stem = csv_path.stem  # e.g. "2025-11"
    accounting_path = output_dir / f"anna_accounting_report_{stem}.xlsx"
    accounting_path = create_accounting_report(df, accounting_path)

    # 生成 Invoice DOCX（按净成本 × 1.15）
    invoice_path = output_dir / f"Invoice_UK_to_HK_{stem}.docx"
    invoice_path = create_invoice_docx(df, csv_path, invoice_path)

    return accounting_path, invoice_path



if __name__ == "__main__":
    import sys

    if len(sys.argv) < 3:
        print("Usage: python anna_monthly_reports_v3.py <csv_path> <output_dir>")
    else:
        acc_path, inv_path = generate_anna_monthly_reports(sys.argv[1], sys.argv[2])
        print(f"[DONE] Accounting: {acc_path}")
        print(f"[DONE] Invoice:    {inv_path}")
