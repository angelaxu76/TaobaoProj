import os
from pathlib import Path
from datetime import date, datetime

import psycopg2
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from config import PGSQL_CONFIG          # 和 manage_export_shipments_costs 一样的数据库配置
from finance_config import FINANCE_EES   # 出口方 / 收货方信息 + 银行信息

# 利润率（UK → HK 批发的加成比例）
MARGIN_RATE = 0.15

# 签名相关（和 anna_monthly_reports_v2 对齐）
SIGN_NAME = "XIAODAN MA"
SIGN_TITLE = "Director, EMINZORA TRADE LTD"
SIGN_IMAGE = r"D:\OneDrive\CrossBorderDocs_UK\00_Templates\signatures\xiaodan_ma_signature.png"


def get_conn():
    return psycopg2.connect(**PGSQL_CONFIG)


def _set_paragraph_style(p, bold=False, size=11, align=None):
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
    if align is not None:
        p.alignment = align


def load_poe_lines(poe_id: str) -> pd.DataFrame:
    """
    从 export_shipments 读取某个 POE 下的所有【商品行】：

    - 只保留 skuid 非空的行（真正有 SKU 的才视为商品）
    - 对这些商品行：
        * 如果有采购成本 -> 用于生成 CI / 内部成本报表
        * 如果没有采购成本 -> 视为非 ANNA 采购 / HK 旧库存，不计入 CI，仅打印警告
    - 那些 skuid 为空的“汇总/物流行”会被跳过，不参与发票金额计算
    """
    sql = """
        SELECT
            id,
            shipment_id,
            skuid,
            product_description,
            hs_code,
            quantity,
            poe_id,
            poe_mrn,
            poe_office,
            poe_date,
            carrier_name,
            tracking_no,
            purchase_unit_cost_gbp
        FROM public.export_shipments
        WHERE poe_id = %s
        ORDER BY shipment_id, skuid, id;
    """
    with get_conn() as conn:
        df = pd.read_sql(sql, conn, params=(poe_id,))

    if df.empty:
        raise ValueError(f"export_shipments 中没有找到 poe_id = {poe_id} 的记录。")

    # 规范化 skuid，过滤出真正的商品行
    df["skuid"] = df["skuid"].fillna("").astype(str).str.strip()
    df_products = df[df["skuid"] != ""].copy()

    if df_products.empty:
        raise ValueError(
            f"poe_id = {poe_id} 没有任何带 SKU 的商品行，"
            f"可能只有一条汇总/物流记录，这种情况无法生成发票。"
        )

    # 按是否填写采购成本拆分
    missing_mask = df_products["purchase_unit_cost_gbp"].isna()
    df_missing = df_products[missing_mask]
    df_with_cost = df_products[~missing_mask].copy()

    # 如果有未填成本的商品行：仅警告，不再阻止生成 CI
    if not df_missing.empty:
        example_skuids = df_missing["skuid"].head(5).tolist()
        print(
            f"[WARN] poe_id = {poe_id} 中有 {len(df_missing)} 条商品未填写采购成本，"
            f"将视为非 ANNA 采购 / HK 旧库存，不计入本次 CI。"
            f" 示例缺失 SKU: {example_skuids}"
        )

    # 如果这一票里完全没有任何有成本的商品，就没必要出 CI 了
    if df_with_cost.empty:
        raise ValueError(
            f"poe_id = {poe_id} 中没有任何填写采购成本的商品，"
            f"推断为全部为 HK 旧库存 / 非 ANNA 采购。"
            f"这种 POE 不需要生成 CI。"
        )

    return df_with_cost


def build_cost_and_price(df: pd.DataFrame) -> pd.DataFrame:
    """
    你在 Excel 填写的是含 VAT 的价格（gross price）。
    这里自动做三件事：

      1）自动去掉 20% VAT 得到净成本 net_cost_unit
      2）在净成本基础上加 15% 批发加成（unit_price_gbp）
      3）计算行金额 line_price_gbp
    """
    df = df.copy()
    df["quantity"] = df["quantity"].fillna(0).astype(int)

    VAT_RATE = 0.20  # 英国标准 VAT 20%

    # 你填入的含税价（gross cost）
    df["gross_cost_unit"] = df["purchase_unit_cost_gbp"].astype(float)

    # 1) 自动转换成净价（net price）
    df["net_cost_unit"] = (df["gross_cost_unit"] / (1 + VAT_RATE)).round(4)

    # 成本小计（用于内部参考，不出现在发票上）
    df["line_net_cost"] = (df["net_cost_unit"] * df["quantity"]).round(2)

    # 2) 批发单价 = 净价 * (1 + 15%)
    df["unit_price_gbp"] = (df["net_cost_unit"] * (1 + MARGIN_RATE)).round(4)

    # 3) 行发票金额 = 批发单价 * 数量
    df["line_price_gbp"] = (df["unit_price_gbp"] * df["quantity"]).round(2)

    return df


def create_poe_cost_report(df: pd.DataFrame, output_path: Path) -> Path:
    """
    生成单个 POE 的内部成本/售价明细 Excel：
      Sheet: POE_CostBreakdown
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    cols = [
        "shipment_id",
        "skuid",
        "product_description",
        "hs_code",
        "quantity",
        "purchase_unit_cost_gbp",
        "line_net_cost",
        "unit_price_gbp",
        "line_price_gbp",
        "poe_id",
        "poe_mrn",
        "poe_office",
        "poe_date",
        "carrier_name",
        "tracking_no",
        "id",
    ]
    df_out = df[cols].copy()

    with pd.ExcelWriter(output_path, engine="xlsxwriter") as writer:
        df_out.to_excel(writer, sheet_name="POE_CostBreakdown", index=False)

    print(f"[OK] POE 成本报表: {output_path}")
    return output_path


def _make_poe_invoice_no(poe_id: str, poe_date: date | None) -> str:
    """
    生成发票号，例如：
      CI-POE-20251008-SD10009875605045
    如果没有 poe_date，则简化为：
      CI-POE-SD10009875605045
    """
    if poe_date:
        ymdd = poe_date.strftime("%Y%m%d")
        return f"CI-POE-{ymdd}-{poe_id}"
    else:
        return f"CI-POE-{poe_id}"


def create_poe_invoice_docx(df: pd.DataFrame, output_path: Path) -> Path:
    """
    根据单个 POE 的行明细生成对香港的商业发票（非 Intercompany）。

    v2 调整：
      - 明细行 No. 重置索引，从 1 连续编号；
      - 去掉 Tracking Number，只在有承运人时显示 Carrier；
      - 不再在正文中写死“(v4.2)”版本号；
      - 增加 Bank Details（从 FINANCE_EES['bank'] 读取）。
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # 汇总信息
    row0 = df.iloc[0]
    poe_id = row0["poe_id"]
    poe_mrn = row0.get("poe_mrn") or ""
    poe_office = row0.get("poe_office") or ""
    poe_date = row0.get("poe_date")
    poe_date_obj = poe_date.date() if isinstance(poe_date, pd.Timestamp) else poe_date

    carrier_name = row0.get("carrier_name") or ""
    tracking_no = row0.get("tracking_no") or ""  # v2: 我们不再打印 tracking_no

    price_total = df["line_price_gbp"].sum().round(2)

    exporter = FINANCE_EES["exporter"]
    consignee = FINANCE_EES["consignee"]
    bank = FINANCE_EES.get("bank", {})
    logistics_cfg = FINANCE_EES.get("logistics", {})

    # 如果数据库中没有 carrier_name，使用配置中的默认（例如 ECMS）
    if not carrier_name:
        carrier_name = logistics_cfg.get("carrier", "") or ""

    invoice_no = _make_poe_invoice_no(poe_id, poe_date_obj)
    invoice_date = (poe_date_obj or datetime.today().date()).isoformat()

    doc = Document()

    # ---------------- Title ----------------
    title = doc.add_heading("COMMERCIAL INVOICE", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------- Header Fields ----------------
    p = doc.add_paragraph(f"Invoice No.: {invoice_no}")
    _set_paragraph_style(p)

    p = doc.add_paragraph(f"Invoice Date: {invoice_date}")
    _set_paragraph_style(p)

    p = doc.add_paragraph(f"Related Proof of Export (POE): {poe_id}")
    _set_paragraph_style(p)

    if poe_mrn:
        p = doc.add_paragraph(f"MRN / Customs Reference: {poe_mrn}")
        _set_paragraph_style(p)

    if poe_office:
        p = doc.add_paragraph(f"Export Office: {poe_office}")
        _set_paragraph_style(p)

    # v2: 不再显示 tracking_no，仅在有承运人时显示 Carrier
    if carrier_name:
        p = doc.add_paragraph(f"Carrier: {carrier_name}")
        _set_paragraph_style(p)

    doc.add_paragraph()

    # ---------------- Seller ----------------
    p = doc.add_paragraph("Seller (UK):\n")
    _set_paragraph_style(p, bold=True)

    p = doc.add_paragraph(
        f"{exporter['name']}\n"
        f"Company No: {exporter.get('company_no','')}\n"
        f"VAT No: {exporter.get('vat_no','')}\n"
        f"EORI: {exporter.get('eori_no','')}\n"
        f"Address: {exporter['address']}\n"
        f"Email: {exporter.get('email','')}\n"
        f"Phone: {exporter.get('phone','')}"
    )
    _set_paragraph_style(p)

    doc.add_paragraph()

    # ---------------- Buyer ----------------
    p = doc.add_paragraph("Buyer (HK):\n")
    _set_paragraph_style(p, bold=True)

    p = doc.add_paragraph(
        f"{consignee['name']}\n"
        f"Address: {consignee['address']}\n"
        f"Email: {consignee.get('email','')}\n"
        f"Phone: {consignee.get('phone','')}"
    )
    _set_paragraph_style(p)

    doc.add_paragraph()

    # ---------------- Terms & Pricing ----------------
    p = doc.add_paragraph("Terms and Pricing:\n")
    _set_paragraph_style(p, bold=True)

    # v2: 去掉 “(v4.2)” 固定版本号，改为更通用的 “the Agreement”
    p = doc.add_paragraph(
        "This Commercial Invoice is issued under the UK–HK Cross-Border Trade Agreement "
        "(the “Agreement”). The goods are sold by the UK Seller to the Hong Kong Buyer "
        "on a wholesale basis.\n\n"
        "Delivery Terms: FCA (ECMS UK warehouse), Incoterms® 2020.\n\n"
        "Pricing Method: The unit prices are determined by applying the agreed cost-plus 15% "
        "wholesale margin to the Seller's net landed cost, in accordance with the agreed Pricing Policy.\n\n"
        "VAT Treatment: This supply qualifies as a zero-rated export of goods for UK VAT "
        "purposes under HMRC Notice 703."
    )
    _set_paragraph_style(p, size=10)

    doc.add_paragraph()

    # ---------------- Item Table ----------------
    # v2: 重置索引，让 No. 从 1 连续编号
    df = df.reset_index(drop=True)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "No."
    hdr[1].text = "SKU / Code"
    hdr[2].text = "Description"
    hdr[3].text = "HS Code"
    hdr[4].text = "Qty"
    hdr[5].text = "Amount (GBP)"

    for idx, row in df.iterrows():
        rc = table.add_row().cells
        rc[0].text = str(idx + 1)
        rc[1].text = str(row["skuid"])
        rc[2].text = str(row.get("product_description") or "")
        rc[3].text = str(row.get("hs_code") or "")
        rc[4].text = str(int(row["quantity"]))
        rc[5].text = f"{row['line_price_gbp']:,.2f}"

    # 合计
    rc = table.add_row().cells
    rc[0].text = ""
    rc[1].text = ""
    rc[2].text = "TOTAL"
    rc[3].text = ""
    rc[4].text = ""
    rc[5].text = f"{price_total:,.2f}"

    doc.add_paragraph()

    # >>> 最终金额行 <<<
    p = doc.add_paragraph()
    p.add_run(f"Total Amount Payable (GBP): {price_total:,.2f}")
    _set_paragraph_style(p, bold=True, size=11)

    doc.add_paragraph()  # blank line

    # ---------------- Bank Details ----------------
    # 从 finance_config.FINANCE_EES["bank"] 读取银行信息
    if bank:
        p = doc.add_paragraph("Bank Details (for settlement):")
        _set_paragraph_style(p, bold=True)

        bank_lines = []
        if bank.get("bank_name"):
            bank_lines.append(f"Bank: {bank['bank_name']}")
        if bank.get("account_name"):
            bank_lines.append(f"Account Name: {bank['account_name']}")
        if bank.get("account_no"):
            bank_lines.append(f"Account No.: {bank['account_no']}")
        if bank.get("sort_code"):
            bank_lines.append(f"Sort Code: {bank['sort_code']}")
        if bank.get("iban"):
            bank_lines.append(f"IBAN: {bank['iban']}")
        if bank.get("swift"):
            bank_lines.append(f"SWIFT / BIC: {bank['swift']}")

        if bank_lines:
            p = doc.add_paragraph("\n".join(bank_lines))
            _set_paragraph_style(p, size=10)

        doc.add_paragraph()  # 空一行和签名区隔开

    # ---------------- Signature ----------------
    p = doc.add_paragraph("Authorised Signature (UK):")
    _set_paragraph_style(p, bold=True)

    if SIGN_IMAGE and os.path.exists(SIGN_IMAGE):
        sign_para = doc.add_paragraph()
        run = sign_para.add_run()
        run.add_picture(SIGN_IMAGE, width=Cm(4))

    p = doc.add_paragraph(f"Name: {SIGN_NAME}")
    _set_paragraph_style(p)
    p = doc.add_paragraph(f"Title: {SIGN_TITLE}")
    _set_paragraph_style(p)
    p = doc.add_paragraph(f"Date: {invoice_date}")
    _set_paragraph_style(p)

    doc.save(output_path)
    print(f"[OK] Commercial Invoice DOCX (no internal reference): {output_path}")
    return output_path


def create_invoice_pdf(docx_path: Path, pdf_path: Path) -> Path:
    try:
        from docx2pdf import convert
    except ImportError:
        print("[WARN] docx2pdf 未安装，跳过 PDF 导出。如需 PDF，请: pip install docx2pdf")
        return pdf_path

    try:
        convert(str(docx_path), str(pdf_path))
        print(f"[OK] Commercial Invoice PDF: {pdf_path}")
    except Exception as e:
        print(f"[WARN] 生成 PDF 失败: {e}")
    return pdf_path


def generate_poe_invoice_and_report(poe_id: str, output_dir: str):
    """
    主入口：按单个 POE 生成：
      1）商业发票（DOCX + 可选 PDF）
      2）POE 成本/售价明细 Excel

    参数：
      poe_id:          e.g. "SD10009905718779"
      output_dir:      e.g. r"D:\OneDrive\CrossBorderDocs_UK\06_Export_Proofs\20251022"
    """
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    df_raw = load_poe_lines(poe_id)
    df = build_cost_and_price(df_raw)

    # 1) Excel 明细
    excel_path = out_dir / f"POE_CostBreakdown_{poe_id}.xlsx"
    excel_path = create_poe_cost_report(df, excel_path)

    # 2) 发票 DOCX + PDF
    invoice_docx_path = out_dir / f"CommercialInvoice_PoE_{poe_id}.docx"
    invoice_docx_path = create_poe_invoice_docx(df, invoice_docx_path)

    invoice_pdf_path = out_dir / f"CommercialInvoice_PoE_{poe_id}.pdf"
    create_invoice_pdf(invoice_docx_path, invoice_pdf_path)

    return excel_path, invoice_docx_path, invoice_pdf_path


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 3:
        print("Usage: python generate_poe_invoice.py <poe_id> <output_dir>")
        sys.exit(1)

    poe_id_arg = sys.argv[1]
    output_dir_arg = sys.argv[2]

    excel_p, docx_p, pdf_p = generate_poe_invoice_and_report(poe_id_arg, output_dir_arg)
    print(f"[DONE] Excel report: {excel_p}")
    print(f"[DONE] Invoice DOCX: {docx_p}")
    print(f"[DONE] Invoice PDF : {pdf_p}")
