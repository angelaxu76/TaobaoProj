from pathlib import Path
from datetime import datetime

from docx import Document
from finance_config import FINANCE_EES


def generate_payment_receipt_docx(
    payment_date: str,
    amount_gbp: float,
    output_path: str,
    closing_balance: float | None = None,
    hk_ref: str | None = None,
    uk_ref: str | None = None,
    receipt_no: str | None = None,
) -> Path:
    """
    生成香港→英国公司付款的 Payment Receipt（收款确认函）DOCX 文件（精简版文案）。

    参数：
        payment_date: 付款日期，格式建议 "YYYY-MM-DD"
        amount_gbp:   收到金额（英镑）
        output_path:  输出文件路径，可以是：
                      1）完整文件路径，如 "D:/.../PaymentReceipt_2025-10-16_10000.docx"
                      2）目录路径，如 "D:/.../04_PaymentProofs"（会自动生成文件名）
        closing_balance: 该笔付款后 intercompany ledger 的余额（可选）
        hk_ref:       香港银行付款参考（可选）
        uk_ref:       英国银行入账参考（可选）
        receipt_no:   回执编号（可选，不传则自动生成）

    返回：
        实际生成的文件 Path 对象
    """

    exporter = FINANCE_EES["exporter"]    # 英国公司信息
    consignee = FINANCE_EES["consignee"]  # 香港公司信息
    bank = FINANCE_EES["bank"]            # 银行信息

    # 解析日期用于显示
    try:
        dt = datetime.strptime(payment_date, "%Y-%m-%d")
        date_str_for_body = dt.strftime("%d %B %Y")  # 16 October 2025
    except ValueError:
        # 如果不是标准 YYYY-MM-DD，就原样使用
        date_str_for_body = payment_date

    # 金额整数部分，用于自动生成编号和文件名
    amount_int = int(round(amount_gbp))

    # 回执编号
    if receipt_no is None:
        date_part = payment_date.replace("-", "")
        receipt_no = f"PR-{date_part}-{amount_int}"

    # 处理 output_path：既支持“目录”，也支持“完整文件名”
    out_path = Path(output_path)
    if out_path.is_dir() or not out_path.suffix:
        # 当作目录用，自动生成文件名
        out_dir = out_path
        out_dir.mkdir(parents=True, exist_ok=True)
        file_name = f"PaymentReceipt_{payment_date}_{amount_int}.docx"
        out_path = out_dir / file_name
    else:
        # 是完整文件路径，确保目录存在
        out_path.parent.mkdir(parents=True, exist_ok=True)

    # 开始生成文档
    doc = Document()

    # 标题
    doc.add_heading("PAYMENT RECEIPT / REMITTANCE CONFIRMATION", level=1)

    # 精简说明：一行
    p = doc.add_paragraph()
    p.add_run("This document confirms that the payment below has been received by the Seller.")

    doc.add_paragraph()  # 空行

    # Receipt No. & 日期
    doc.add_paragraph(f"Receipt No.: {receipt_no}")
    doc.add_paragraph(f"Receipt Date: {date_str_for_body}")

    doc.add_paragraph()

    # Seller 信息（英国公司）
    doc.add_paragraph("Seller (UK Company):")
    doc.add_paragraph(f"  {exporter['name']}")
    doc.add_paragraph(f"  {exporter['address']}")
    if exporter.get("company_no"):
        doc.add_paragraph(f"  Company No: {exporter['company_no']}")
    if exporter.get("vat_no"):
        doc.add_paragraph(f"  VAT No: {exporter['vat_no']}")
    if exporter.get("eori_no"):
        doc.add_paragraph(f"  EORI No: {exporter['eori_no']}")

    doc.add_paragraph()

    # Buyer 信息（香港公司）
    doc.add_paragraph("Buyer (Hong Kong Company):")
    doc.add_paragraph(f"  {consignee['name']}")
    doc.add_paragraph(f"  {consignee['address']}")

    doc.add_paragraph()

    # 分隔线 + 付款详情
    doc.add_paragraph("-" * 66)
    doc.add_paragraph("PAYMENT DETAILS")
    doc.add_paragraph("-" * 66)

    doc.add_paragraph(f"• Payment Date: {date_str_for_body}")
    doc.add_paragraph(f"• Amount Received (GBP): {amount_gbp:,.2f}")

    if hk_ref:
        doc.add_paragraph(f"• HK Payer Bank Reference No.: {hk_ref}")
    else:
        doc.add_paragraph("• HK Payer Bank Reference No.: ____________________")

    doc.add_paragraph(
        f"• UK Receiving Bank: {bank.get('bank_name', 'ANNA Money')}"
    )
    if bank.get("account_name"):
        doc.add_paragraph(f"• UK Account Name: {bank['account_name']}")
    if bank.get("account_no"):
        doc.add_paragraph(f"• UK Account No.: {bank['account_no']}")
    if bank.get("sort_code"):
        doc.add_paragraph(f"• Sort Code: {bank['sort_code']}")

    if uk_ref:
        doc.add_paragraph(f"• UK Bank Reference / Faster Payment ID: {uk_ref}")
    else:
        doc.add_paragraph(
            "• UK Bank Reference / Faster Payment ID: ____________________"
        )

    doc.add_paragraph()

    # 应用到 intercompany balance —— 精简两句
    doc.add_paragraph("-" * 66)
    doc.add_paragraph("APPLICATION TO INTERCOMPANY TRADE BALANCE")
    doc.add_paragraph("-" * 66)

    doc.add_paragraph(
        "This payment has been recorded in the Intercompany Ledger as settlement for goods supplied under the cost-plus resale model."
    )

    doc.add_paragraph(
        "This receipt is not a commercial invoice; the official Commercial Invoice (CI) for the billing period is issued separately."
    )

    doc.add_paragraph()
    doc.add_paragraph("Ledger Entry:")
    doc.add_paragraph(f"• Debit (Cash Received): GBP {amount_gbp:,.2f}")
    doc.add_paragraph(f"• Credit (Trade Receivable Reduction): GBP {amount_gbp:,.2f}")

    if closing_balance is not None:
        doc.add_paragraph(
            f"• Updated Intercompany Balance: GBP {closing_balance:,.2f}"
        )
    else:
        doc.add_paragraph("• Updated Intercompany Balance: GBP __________")

    doc.add_paragraph()

    # DECLARATION —— 精简为一段
    doc.add_paragraph("-" * 66)
    doc.add_paragraph("DECLARATION")
    doc.add_paragraph("-" * 66)

    doc.add_paragraph(
        "The above payment has been recorded as part of the intercompany settlement. Supporting bank statements are kept and can be provided on request."
    )

    doc.add_paragraph()
    doc.add_paragraph("Authorized by:")
    doc.add_paragraph("Name: ____________________")
    doc.add_paragraph("Title: ____________________")
    doc.add_paragraph("Date: ____________________")

    # 保存
    doc.save(out_path)
    print(f"[OK] Payment receipt generated: {out_path}")
    return out_path
