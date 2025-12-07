import csv
from datetime import datetime
from pathlib import Path
import os

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

# finance_config.py 与本文件在同一目录
from finance_config import FINANCE_EES


# ✅ 不向香港公司收费的类别（英国自己承担的 or 不相关的）
EXCLUDE_CATEGORIES = {
    "Sales",                 # 香港打款，不是成本/收入
    "Accountant",            # 会计费用：英国公司承担
    "Business account fees", # ANNA 账户费：英国公司承担
}

# ✅ 董事自动电子签名配置
SIGN_NAME = "XIAODAN MA"  # 如需改为 Nianzhou，只改这几行
SIGN_TITLE = "Director, EMINZORA TRADE LTD"
SIGN_IMAGE = r"D:\OneDrive\CrossBorderDocs_UK\00_Templates\signatures\xiaodan_ma_signature.png"
# 如果没有签名图片，脚本会自动退回为下划线签名占位


def load_and_prepare(csv_path: str) -> pd.DataFrame:
    """
    读取 ANNA 导出的 CSV，做基础清洗，并返回 DataFrame.

    关键点：
    - 容忍坏行（列数不对），自动截断/补齐
    - 转换 Amount 为 float
    - 解析 Created 为 datetime
    - 添加 Item_Type 分类列
    - 过滤掉 EXCLUDE_CATEGORIES（不向香港收费）
    - 添加 Net_Ex_VAT / VAT_Amount（按 20% 反推）
    """
    csv_path = Path(csv_path)
    print(f"[INFO] Loading CSV: {csv_path}")

    rows = []
    with csv_path.open("r", encoding="utf-8", newline="") as f:
        reader = csv.reader(f)
        try:
            header = next(reader)
        except StopIteration:
            raise ValueError("Empty CSV")

        expected_cols = len(header)

        for line_no, row in enumerate(reader, start=2):
            if not any(row):
                # 空行跳过
                continue

            if len(row) < expected_cols:
                # 列数不够 -> 用空字符串补齐
                row = row + [""] * (expected_cols - len(row))
            elif len(row) > expected_cols:
                # 列数太多 -> 截断
                row = row[:expected_cols]

            rows.append(row)

    df = pd.DataFrame(rows, columns=header)
    # 去掉全空列
    df = df.dropna(axis=1, how="all")

    # ---- 解析日期 ----
    if "Created" in df.columns:
        # ANNA 一般是 "2025-11-02 10:23:45"
        df["Created_dt"] = pd.to_datetime(
            df["Created"].str.replace(",", ""),
            format="%Y-%m-%d %H:%M:%S",
            errors="coerce",
        )
    else:
        df["Created_dt"] = pd.NaT

    # ---- 金额转 float ----
    if "Amount" not in df.columns:
        raise ValueError("CSV 中缺少 Amount 列")

    df["Amount"] = pd.to_numeric(df["Amount"], errors="coerce").fillna(0.0)

    # ---- 分类（仅用于展示）----
    def classify_item_type(category: str, description: str) -> str:
        category = (category or "").lower()
        description = (description or "").lower()

        if "stock" in category or "inventory" in category:
            return "Stock / Goods"
        if "postage" in category or "shipping" in category or "delivery" in category:
            return "Shipping / Postage"
        if "materials" in category or "packag" in category:
            return "Packaging / Materials"
        if "parcel" in category or "courier" in description:
            return "Courier / Parcel"
        if "refund" in category or "return" in description:
            return "Refund / Return"
        if "fee" in category:
            return "Fees / Charges"
        return "Other cost"

    df["Item_Type"] = df.apply(
        lambda r: classify_item_type(
            r.get("Category", "") if "Category" in df.columns else "",
            r.get("Description", "") if "Description" in df.columns else "",
        ),
        axis=1,
    )

    # ---- 过滤掉不向香港收费的类别 ----
    if "Category" in df.columns:
        before = len(df)
        df = df[~df["Category"].isin(EXCLUDE_CATEGORIES)].copy()
        after = len(df)
        print(f"[INFO] Filter EXCLUDE_CATEGORIES: {before} -> {after} rows")

    # ---- 计算 Net_Ex_VAT / VAT_Amount（按 20% 反推） ----
    # 约定：Amount 为负数代表支出，为正数代表退款
    # 所以直接 Amount / 1.2 即为“含符号的净额”
    df["Net_Ex_VAT"] = (df["Amount"] / 1.2).round(2)
    df["VAT_Amount"] = (df["Amount"] - df["Net_Ex_VAT"]).round(2)

    return df


def create_accounting_report(df: pd.DataFrame, output_path: Path) -> Path:
    """
    创建一个 Excel 记账报表：
    - Sheet1: 明细（含 Amount, Net_Ex_VAT, VAT_Amount, Item_Type）
    - Sheet2: 汇总（按 Item_Type 分组 + 总计）
    - 再附上 Net total / VAT total / Gross total / UK Profit / HK Payable

    注意：这个 Excel 是“后台成本计算 / 支撑文件”，不是给香港公司的正式发票。
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    print(f"[INFO] Writing accounting report to: {output_path}")

    # 只保留一些对你记账有意义的字段
    detail_cols = []
    for col in [
        "Created",
        "Created_dt",
        "Description",
        "Category",
        "Amount",
        "Net_Ex_VAT",
        "VAT_Amount",
        "Item_Type",
        "Reference",
        "Counterparty name",
    ]:
        if col in df.columns:
            detail_cols.append(col)

    df_detail = df[detail_cols].copy()

    # 计算总额
    gross_total = df["Amount"].sum().round(2)          # 含 VAT（带符号）
    net_total = df["Net_Ex_VAT"].sum().round(2)        # 去 VAT（带符号）
    vat_total = df["VAT_Amount"].sum().round(2)

    # 取“正向成本数值”
    gross_cost = round(-gross_total, 2)  # 负数 → 成本
    net_cost = round(-net_total, 2)

    uk_profit = round(net_cost * 0.15, 2)
    hk_payable = round(net_cost * 1.15, 2)

    # 分组汇总
    grp = df.groupby("Item_Type")[["Amount", "Net_Ex_VAT", "VAT_Amount"]].sum().round(2)
    grp = grp.reset_index()

    # 写 Excel
    with pd.ExcelWriter(output_path, engine="xlsxwriter") as writer:
        df_detail.to_excel(writer, sheet_name="Details", index=False)

        summary_rows = []

        for _, row in grp.iterrows():
            summary_rows.append(
                {
                    "Item_Type": row["Item_Type"],
                    "Amount (Gross)": row["Amount"],
                    "Net_Ex_VAT": row["Net_Ex_VAT"],
                    "VAT_Amount": row["VAT_Amount"],
                }
            )

        summary_rows.append({})
        summary_rows.append({"Item_Type": "TOTAL (signed)", "Amount (Gross)": gross_total,
                             "Net_Ex_VAT": net_total, "VAT_Amount": vat_total})
        summary_rows.append({"Item_Type": "COST (positive)", "Amount (Gross)": gross_cost,
                             "Net_Ex_VAT": net_cost, "VAT_Amount": vat_total})
        summary_rows.append({"Item_Type": "UK Profit (15%)", "Amount (Gross)": uk_profit})
        summary_rows.append({"Item_Type": "HK Payable (Net*1.15)", "Amount (Gross)": hk_payable})

        df_summary = pd.DataFrame(summary_rows)
        df_summary.to_excel(writer, sheet_name="Summary", index=False)

    print(f"[OK] Accounting report: {output_path}")
    return output_path


def _set_paragraph_style(p, bold=False, size=11, align=None):
    """小工具：设置段落基本样式。"""
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
    if align is not None:
        p.alignment = align


def _make_invoice_number(start_date, csv_path: Path) -> str:
    """
    生成一个唯一的发票编号，例如：
    INV-UKHK-202511
    """
    try:
        if start_date:
            ym = f"{start_date.year}{start_date.month:02d}"
            return f"INV-UKHK-{ym}"
        # fallback: 从文件名推日期，如 2025-11.csv
        stem = csv_path.stem
        dt = datetime.strptime(stem, "%Y-%m")
        ym = f"{dt.year}{dt.month:02d}"
        return f"INV-UKHK-{ym}"
    except Exception:
        return "INV-UKHK-UNKNOWN"


def create_invoice_docx(
    df: pd.DataFrame,
    csv_path: Path,
    output_path: Path,
) -> Path:
    """
    根据 df 生成一个正式的 Intercompany Commercial Invoice（DOCX，英文商务版）。

    金额逻辑：
    - gross_total = df["Amount"].sum()        （带符号，负数为成本）
    - net_total   = df["Net_Ex_VAT"].sum()
    - gross_cost  = -gross_total             （正数：含 VAT 成本）
    - net_cost    = -net_total               （正数：去 VAT 成本 = HK 采购基数）
    - uk_profit   = net_cost * 0.15
    - hk_payable  = net_cost * 1.15          （发票金额，Cost Plus 15%）
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    print(f"[INFO] Writing invoice DOCX to: {output_path}")

    gross_total = df["Amount"].sum().round(2)
    net_total = df["Net_Ex_VAT"].sum().round(2)
    vat_total = df["VAT_Amount"].sum().round(2)

    gross_cost = round(-gross_total, 2)
    net_cost = round(-net_total, 2)
    uk_profit = round(net_cost * 0.15, 2)
    hk_payable = round(net_cost * 1.15, 2)

    # 取账期：按 Created_dt 的最小/最大日期
    if df["Created_dt"].notna().any():
        start_date = df["Created_dt"].min().date()
        end_date = df["Created_dt"].max().date()
    else:
        # fallback：从文件名推一个月份
        try:
            # e.g. 2025-11.csv
            stem = Path(csv_path).stem
            dt = datetime.strptime(stem, "%Y-%m")
            start_date = dt.date().replace(day=1)
            end_date = dt.date()
        except Exception:
            start_date = None
            end_date = None

    invoice_no = _make_invoice_number(start_date, csv_path)
    exporter = FINANCE_EES["exporter"]
    consignee = FINANCE_EES["consignee"]

    doc = Document()

    # 标题：COMMERCIAL INVOICE
    title = doc.add_heading("COMMERCIAL INVOICE", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 副标题：UK → HK Intercompany Supply
    sub = doc.add_paragraph("UK → HK Intercompany Supply (Cost-Plus 15% Wholesale Pricing)")
    _set_paragraph_style(sub, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 发票号
    p = doc.add_paragraph()
    p.add_run(f"Invoice No.: {invoice_no}")
    _set_paragraph_style(p, size=11)

    # 发票日期（自动今天）
    invoice_date = datetime.today().date().isoformat()
    p = doc.add_paragraph()
    p.add_run(f"Invoice Date: {invoice_date}")
    _set_paragraph_style(p, size=11)

    # 账期
    if start_date and end_date:
        period_text = f"{start_date.isoformat()} to {end_date.isoformat()}"
    else:
        period_text = "See attached ANNA statement"

    p = doc.add_paragraph()
    p.add_run(f"Period Covered: {period_text}")
    _set_paragraph_style(p, size=11)

    doc.add_paragraph()  # 空行

    # 出票方（英国公司）
    p = doc.add_paragraph()
    p.add_run("Seller (UK):\n")
    _set_paragraph_style(p, bold=True, size=11)
    p.add_run(f"{exporter['name']}\n")
    p.add_run(f"Company No: {exporter.get('company_no', '')}\n")
    p.add_run(f"VAT No: {exporter.get('vat_no', '')}\n")
    p.add_run(f"EORI: {exporter.get('eori_no', '')}\n")
    p.add_run(f"Address: {exporter['address']}\n")
    p.add_run(f"Email: {exporter.get('email', '')}\n")
    p.add_run(f"Phone: {exporter.get('phone', '')}")

    doc.add_paragraph()  # 空行

    # 收票方（香港公司）
    p = doc.add_paragraph()
    p.add_run("Buyer (HK):\n")
    _set_paragraph_style(p, bold=True, size=11)
    p.add_run(f"{consignee['name']}\n")
    p.add_run(f"Address: {consignee['address']}\n")
    p.add_run(f"Email: {consignee.get('email', '')}\n")
    p.add_run(f"Phone: {consignee.get('phone', '')}")

    doc.add_paragraph()  # 空行

    # 定价说明
    p = doc.add_paragraph()
    p.add_run("Pricing Basis:\n")
    _set_paragraph_style(p, bold=True, size=11)
    p.add_run(
        "This Commercial Invoice is issued in accordance with the UK–HK Cross-Border "
        "Trade Agreement (v4). All goods supplied during this period are sold by the "
        "UK Seller to the Hong Kong Buyer under a Cost-Plus 15% wholesale pricing "
        "method. The aggregate invoice amount is calculated as the UK landed net "
        "cost of goods plus a 15% margin."
    )

    doc.add_paragraph()  # 空行

    # 金额汇总表：更标准的商业格式
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Description"
    hdr_cells[1].text = "Amount (GBP)"

    def add_row(label: str, value: float):
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = f"{value:,.2f}"

    add_row("Supply of goods (net of UK VAT)", net_cost)
    add_row("Subtotal", net_cost)
    add_row("Add: agreed wholesale margin (15%)", uk_profit)
    add_row("Total amount due (zero-rated export)", hk_payable)

    doc.add_paragraph()  # 空行

    # Internal calculation reference（只作信息参考，金额转成正数）
    p = doc.add_paragraph()
    p.add_run("Internal calculation reference (for information only):\n")
    _set_paragraph_style(p, bold=False, size=9)
    p.add_run(
        f"- Underlying ANNA gross cost (including UK VAT): {abs(gross_total):,.2f} GBP\n"
        f"- Underlying ANNA net cost (excluding UK VAT): {abs(net_total):,.2f} GBP\n"
        f"- Total UK input VAT on these purchases: {abs(vat_total):,.2f} GBP\n"
        "  (These figures are shown for reconciliation purposes only and do not "
        "represent additional charges to the Buyer. UK input VAT is claimed and "
        "retained by the UK Seller as part of its own VAT return.)"
    )

    doc.add_paragraph()  # 空行

    # 备注
    notes = doc.add_paragraph()
    notes.add_run("Notes:\n")
    _set_paragraph_style(notes, bold=True, size=10)
    notes = doc.add_paragraph()
    notes.add_run(
        f"- This Commercial Invoice (Invoice No.: {invoice_no}) summarises the wholesale "
        "supply of goods from the UK Seller to the Hong Kong Buyer for the above period, "
        "under FCA (ECMS UK warehouse) terms.\n"
        "- Individual export consignments are documented in separate Commercial Invoices (CI) and "
        "Export Evidence Summaries (EES), together with customs Proof of Export (POE) documents. "
        "These are retained by the UK Seller for VAT zero-rating and audit purposes.\n"
        "- The transaction is treated as a zero-rated export of goods for UK VAT purposes, "
        "in line with HMRC Notice 703.\n"
        "- Payment terms: as per the UK–HK Cross-Border Trade Agreement (e.g. payment within 7 days "
        "from invoice date)."
    )
    _set_paragraph_style(notes, size=10)

    doc.add_paragraph()  # 空行

    # 签名区（自动日期 + 电子签名图片）
    sig_para = doc.add_paragraph()
    sig_para.add_run("\nAuthorised Signature (UK):")
    _set_paragraph_style(sig_para, bold=True, size=11)

    # 插入电子签名图片或下划线
    if SIGN_IMAGE and os.path.exists(SIGN_IMAGE):
        pic_para = doc.add_paragraph()
        run = pic_para.add_run()
        run.add_picture(SIGN_IMAGE, width=Cm(4))
    else:
        underline_para = doc.add_paragraph("_______________________________")
        _set_paragraph_style(underline_para, size=11)

    # 签名人信息 + 日期（自动今天）
    name_para = doc.add_paragraph(f"Name: {SIGN_NAME}")
    _set_paragraph_style(name_para, size=10)
    title_para = doc.add_paragraph(f"Title: {SIGN_TITLE}")
    _set_paragraph_style(title_para, size=10)
    date_para = doc.add_paragraph(f"Date: {invoice_date}")
    _set_paragraph_style(date_para, size=10)

    doc.save(output_path)
    print(f"[OK] Commercial Invoice DOCX: {output_path}")
    return output_path


def generate_anna_monthly_reports(csv_path: str, output_dir: str):
    """
    主入口函数（供 pipeline 调用）：

    输入：
        csv_path: ANNA 月度交易 CSV 路径（例如 2025-11.csv）
        output_dir: 输出目录

    输出：
        (accounting_path, invoice_path)

    逻辑：
        1）生成后台会计报表 anna_accounting_report_YYYY-MM.xlsx
        2）生成 UK→HK 的月度 COMMERCIAL INVOICE（Cost+15% 模式）
    """
    csv_path = Path(csv_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    df = load_and_prepare(str(csv_path))

    # 生成记账报表（后台支撑文件）
    stem = csv_path.stem  # e.g. "2025-11"
    accounting_path = output_dir / f"anna_accounting_report_{stem}.xlsx"
    accounting_path = create_accounting_report(df, accounting_path)

    # 生成 Intercompany Commercial Invoice DOCX（Cost+15%）
    invoice_path = output_dir / f"Invoice_UK_to_HK_{stem}.docx"
    invoice_path = create_invoice_docx(df, csv_path, invoice_path)

    return accounting_path, invoice_path


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 3:
        print("Usage: python anna_monthly_reports_v4.py <csv_path> <output_dir>")
    else:
        acc_path, inv_path = generate_anna_monthly_reports(sys.argv[1], sys.argv[2])
        print(f"[DONE] Accounting: {acc_path}")
        print(f"[DONE] Invoice:    {inv_path}")
